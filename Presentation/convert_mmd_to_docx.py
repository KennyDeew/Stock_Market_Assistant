#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Скрипт для конвертации Mermaid диаграмм в DOCX формат с визуализацией схем
Требует: python-docx, requests, Pillow
"""

import re
import os
import base64
import requests
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from PIL import Image

def add_heading(doc, text, level):
    """Добавить заголовок"""
    heading = doc.add_heading(text, level=level)
    return heading

def render_mermaid_to_image(mermaid_code):
    """
    Рендеринг Mermaid диаграммы в изображение через mermaid.ink API
    Альтернативно можно использовать локальный mermaid-cli или playwright
    """
    try:
        # Используем mermaid.ink API для рендеринга
        # Кодируем диаграмму в base64 URL-safe
        encoded = base64.urlsafe_b64encode(mermaid_code.encode('utf-8')).decode('utf-8').rstrip('=')
        url = f"https://mermaid.ink/img/{encoded}"

        # Загружаем изображение
        response = requests.get(url, timeout=30)
        if response.status_code == 200:
            return BytesIO(response.content)
        else:
            print(f"  Ошибка при рендеринге диаграммы: HTTP {response.status_code}")
            return None
    except ImportError:
        print("  Предупреждение: библиотека requests не установлена. Установите: pip install requests")
        return None
    except Exception as e:
        print(f"  Ошибка при рендеринге Mermaid: {e}")
        return None

def add_image_to_doc(doc, image_stream, width_inches=6.5):
    """Добавить изображение в документ"""
    if image_stream is None:
        return

    try:
        # Позиционируем параграф по центру
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Добавляем изображение
        run = para.add_run()
        image_stream.seek(0)
        run.add_picture(image_stream, width=Inches(width_inches))

        # Добавляем небольшой отступ после изображения
        para.paragraph_format.space_after = Pt(12)

    except Exception as e:
        print(f"Ошибка при добавлении изображения: {e}")
        # Добавляем текст-заглушку
        para = doc.add_paragraph()
        para.add_run(f"[Ошибка загрузки изображения: {e}]").font.size = Pt(9)

def add_code_block_fallback(doc, code, language='mermaid'):
    """Добавить блок кода как запасной вариант"""
    para = doc.add_paragraph()
    para.style = 'No Spacing'

    run = para.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0, 0, 0)

    para_format = para.paragraph_format
    para_format.left_indent = Inches(0.5)
    para_format.right_indent = Inches(0.5)
    para_format.space_before = Pt(6)
    para_format.space_after = Pt(6)

def add_bullet_list(doc, items):
    """Добавить маркированный список"""
    for item in items:
        para = doc.add_paragraph(item, style='List Bullet')
    return doc

def parse_markdown_to_docx(mmd_file, docx_file):
    """Парсинг Markdown файла и создание DOCX документа с визуализированными диаграммами"""

    # Читаем исходный файл
    with open(mmd_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # Создаем новый DOCX документ
    doc = Document()

    # Настройка стилей документа
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Добавляем заголовок документа
    title = doc.add_heading('Stock Market Assistant - Архитектура системы', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Парсим содержимое
    lines = content.split('\n')
    i = 0
    diagram_count = 0

    while i < len(lines):
        line = lines[i].strip()

        # Обработка заголовков
        if line.startswith('# '):
            text = line[2:].strip()
            add_heading(doc, text, 1)
        elif line.startswith('## '):
            text = line[3:].strip()
            add_heading(doc, text, 2)
        elif line.startswith('### '):
            text = line[4:].strip()
            add_heading(doc, text, 3)

        # Обработка блоков кода Mermaid - РЕНДЕРИМ В ИЗОБРАЖЕНИЯ
        elif line.startswith('```mermaid'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            code = '\n'.join(code_lines)

            print(f"Рендеринг диаграммы {diagram_count + 1}...")
            diagram_count += 1

            # Рендерим диаграмму в изображение
            image_stream = render_mermaid_to_image(code)

            if image_stream:
                add_image_to_doc(doc, image_stream, width_inches=6.5)
                print(f"  ✓ Диаграмма {diagram_count} успешно добавлена")
            else:
                # Запасной вариант - добавляем код
                print(f"  ⚠ Не удалось отрендерить диаграмму {diagram_count}, добавляю код")
                add_code_block_fallback(doc, code, 'mermaid')
                note_para = doc.add_paragraph()
                note_run = note_para.add_run('Примечание: Не удалось отрендерить диаграмму. Используйте Mermaid Live Editor (https://mermaid.live) для визуализации.')
                note_run.font.size = Pt(9)
                note_run.font.italic = True
                note_run.font.color.rgb = RGBColor(128, 128, 128)

        # Обработка обычных блоков кода
        elif line.startswith('```'):
            code_lines = []
            language = line[3:].strip() if len(line) > 3 else ''
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            code = '\n'.join(code_lines)
            add_code_block_fallback(doc, code, language)

        # Обработка маркированных списков
        elif line.startswith('- '):
            items = []
            while i < len(lines) and lines[i].strip().startswith('- '):
                items.append(lines[i].strip()[2:])
                i += 1
            add_bullet_list(doc, items)
            i -= 1  # Откатываем на одну строку, так как цикл увеличит i

        # Обработка обычного текста
        elif line and not line.startswith('#'):
            para = doc.add_paragraph(line)

        i += 1

    # Сохраняем документ
    doc.save(docx_file)
    print(f"\n✓ Документ успешно создан: {docx_file}")
    print(f"  Всего диаграмм обработано: {diagram_count}")

if __name__ == '__main__':
    import sys

    # Определяем пути относительно скрипта
    script_dir = os.path.dirname(os.path.abspath(__file__))
    mmd_file = os.path.join(script_dir, 'Архитектура.mmd')
    docx_file = os.path.join(script_dir, 'Архитектура.docx')

    if not os.path.exists(mmd_file):
        print(f"Ошибка: файл {mmd_file} не найден!")
        sys.exit(1)

    print(f"Конвертация {mmd_file} -> {docx_file}")
    print("=" * 50)

    parse_markdown_to_docx(mmd_file, docx_file)

